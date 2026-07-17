"""DOCX resume parser."""

from docx import Document


def parse_docx(file_path: str) -> str:
    """
    Extract raw text from a DOCX resume.

    Args:
        file_path: Path to the DOCX resume.

    Returns:
        Extracted text as a single string.
    """

    document = Document(file_path)

    text = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)

    # Extract tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)